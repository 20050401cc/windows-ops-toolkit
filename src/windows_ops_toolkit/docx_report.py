"""Generate a simple DOCX report from a JSON outline."""

from __future__ import annotations

import argparse
import json
from pathlib import Path


def build_report(data: dict, output: Path) -> Path:
    from docx import Document

    document = Document()
    document.add_heading(data.get("title", "Project Report"), 0)

    subtitle = data.get("subtitle")
    if subtitle:
        document.add_paragraph(subtitle)

    for section in data.get("sections", []):
        document.add_heading(section.get("title", "Section"), level=1)
        for paragraph in section.get("paragraphs", []):
            document.add_paragraph(paragraph)
        rows = section.get("table")
        if rows:
            table = document.add_table(rows=1, cols=len(rows[0]))
            table.style = "Table Grid"
            for index, value in enumerate(rows[0]):
                table.rows[0].cells[index].text = str(value)
            for row in rows[1:]:
                cells = table.add_row().cells
                for index, value in enumerate(row):
                    cells[index].text = str(value)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate a DOCX report from JSON.")
    parser.add_argument("--input", required=True, type=Path, help="Report JSON file.")
    parser.add_argument("--output", required=True, type=Path, help="Output DOCX file.")
    args = parser.parse_args()
    data = json.loads(args.input.read_text(encoding="utf-8"))
    print(build_report(data, args.output))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
