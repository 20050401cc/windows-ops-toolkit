"""Check whether DOCX files are structurally readable."""

from __future__ import annotations

import argparse
import json
import zipfile
from pathlib import Path


def check_docx(path: Path) -> dict:
    result = {
        "path": str(path),
        "exists": path.exists(),
        "ok": False,
        "zip_error": None,
        "paragraph_count": None,
        "table_count": None,
        "image_count": None,
    }
    if not path.exists():
        result["zip_error"] = "file does not exist"
        return result

    try:
        with zipfile.ZipFile(path) as archive:
            result["zip_error"] = archive.testzip()
            names = archive.namelist()
            result["image_count"] = len(
                [name for name in names if name.startswith("word/media/")]
            )
    except zipfile.BadZipFile as exc:
        result["zip_error"] = str(exc)
        return result

    try:
        from docx import Document

        document = Document(path)
        result["paragraph_count"] = len(document.paragraphs)
        result["table_count"] = len(document.tables)
    except Exception as exc:
        result["zip_error"] = f"python-docx read failed: {exc}"
        return result

    result["ok"] = result["zip_error"] is None
    return result


def main() -> int:
    parser = argparse.ArgumentParser(description="Check DOCX integrity.")
    parser.add_argument("files", nargs="+", type=Path)
    parser.add_argument("--json", action="store_true", help="Print JSON output.")
    args = parser.parse_args()

    results = [check_docx(path) for path in args.files]
    if args.json:
        print(json.dumps(results, indent=2, ensure_ascii=False))
    else:
        for item in results:
            status = "OK" if item["ok"] else "FAIL"
            print(
                f"{status} {item['path']} "
                f"paragraphs={item['paragraph_count']} "
                f"tables={item['table_count']} "
                f"images={item['image_count']}"
            )
    return 0 if all(item["ok"] for item in results) else 1


if __name__ == "__main__":
    raise SystemExit(main())
